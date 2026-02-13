import os
import shutil
import tempfile
import time
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
import docx
import pymupdf4llm
from google import genai
from dotenv import load_dotenv

# --- AYARLAR ---
env_path = Path(__file__).parent / ".env"
load_dotenv(dotenv_path=env_path)

API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    raise RuntimeError("KRİTİK HATA: .env dosyasında GEMINI_API_KEY bulunamadı.")

client = genai.Client(api_key=API_KEY)

# YEDEK MODEL (Senin istediğin 2.0 Lite)
FALLBACK_MODEL = "gemini-2.0-flash-lite"

app = FastAPI(title="Next-Gen APA Generator")

def extract_markdown_from_pdf(file_path: str) -> str:
    """PDF'ten metni Markdown olarak çeker."""
    try:
        return pymupdf4llm.to_markdown(file_path, pages=[0, 1])
    except Exception as e:
        print(f"PDF Okuma Hatası: {e}")
        return ""

def generate_citation_with_retry(combined_text: str, model_id: str, retries=2) -> str:
    """
    İstenilen modelle dener. Hata alırsa 2.0 Flash-Lite modeline düşer.
    """
    prompt = (
        "Create a bibliography entry for each of the following academic texts in APA 7 format. "
        "Output ONLY the list of citations. No intro/outro text. One citation per line.\n\n"
        f"{combined_text}"
    )

    print(f"🤖 İstek Gönderiliyor: {model_id}")

    for attempt in range(retries):
        try:
            response = client.models.generate_content(
                model=model_id, 
                contents=prompt
            )
            return response.text.strip()
        
        except Exception as e:
            error_msg = str(e)
            print(f"⚠️ Hata ({model_id}): {error_msg}")

            # Eğer model bulunamazsa veya erişim yoksa (404 / 403)
            if "404" in error_msg or "not found" in error_msg.lower() or "403" in error_msg:
                if model_id != FALLBACK_MODEL:
                    print(f"🔄 Model bulunamadı. '{FALLBACK_MODEL}' modeline geçiliyor...")
                    model_id = FALLBACK_MODEL
                    continue 
            
            # Kota dolduysa (429)
            elif "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                wait_time = 10
                print(f"⏳ Kota sınırına takıldık. {wait_time}sn bekleniyor...")
                time.sleep(wait_time)
                # Yine de çözülmezse yedeğe geçmeyi deneyelim
                if attempt == retries - 1 and model_id != FALLBACK_MODEL:
                     print(f"🔄 Kota açılmadı, daha hafif olan '{FALLBACK_MODEL}' deneniyor.")
                     model_id = FALLBACK_MODEL
                     # Son bir şans daha ver döngüye
                     continue
            
            else:
                # Beklenmeyen kritik hata
                raise HTTPException(status_code=500, detail=f"Yapay Zeka Hatası: {error_msg}")
    
    raise HTTPException(status_code=429, detail="Servis şu an yanıt veremiyor, lütfen daha sonra deneyin.")

def cleanup_temp_file(path: str):
    """Dosya gönderildikten sonra siler."""
    try:
        if os.path.exists(path):
            os.remove(path)
            print(f"🧹 Temizlik yapıldı: {path}")
    except Exception as e:
        print(f"Temizlik uyarısı: {e}")

@app.post("/generate-bibliography/")
async def generate_bibliography(
    background_tasks: BackgroundTasks,
    files: list[UploadFile] = File(...), 
    model_id: str = Form(...)
):
    # Klasör kontrolü bizde (Python otomatik silmesin diye mkdtemp)
    temp_dir = tempfile.mkdtemp()
    
    try:
        combined_prompt_text = ""
        
        # 1. Dosyaları İşle
        for file in files:
            temp_file_path = os.path.join(temp_dir, file.filename)
            with open(temp_file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            text = extract_markdown_from_pdf(temp_file_path)
            if text:
                combined_prompt_text += f"--- KAYNAK: {file.filename} ---\n{text[:5000]}\n\n"

        if not combined_prompt_text:
            raise HTTPException(status_code=400, detail="PDF dosyalarından metin okunamadı.")

        # 2. Yapay Zeka ile Üret
        citations = generate_citation_with_retry(combined_prompt_text, model_id)

        # 3. Word Dosyasını Oluştur
        doc = docx.Document()
        doc.add_heading(f'Kaynakça (APA 7)', 0)
        doc.add_paragraph(citations)
        doc.add_paragraph(f"\n(Oluşturulma Modeli: {model_id})") # Bilgi notu
        
        # Dosyayı temp klasörün dışına güvenli bir yere alalım
        fd, output_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd) 
        
        doc.save(output_path)

        # 4. Temizlik ve Gönderim
        shutil.rmtree(temp_dir) 
        background_tasks.add_task(cleanup_temp_file, output_path)

        return FileResponse(
            path=output_path, 
            filename="kaynakca_apa7.docx", 
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise e

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8001)